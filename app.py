# import streamlit as st
# import google.generativeai as genai

# # -----------------------
# # 🔑 Configure Gemini API
# # -----------------------
# api_key = st.secrets["GOOGLE_API_KEY"]
# genai.configure(api_key=api_key)

# # model = genai.GenerativeModel("gemini-1.5-flash")
# model = genai.GenerativeModel("gemini-2.5-flash")

# # -----------------------
# # 🌟 Page Config
# # -----------------------
# st.set_page_config(
#     page_title="AI Resume Reviewer",
#     page_icon="📄",
#     layout="wide"
# )

# # -----------------------
# # 🎯 Header Section
# # -----------------------
# st.markdown("""
# <h1 style='text-align:center; margin-bottom:0;'>📄 AI Resume Reviewer</h1>
# <p style='text-align:center; font-size:18px; color:gray;'>
# Upload your resume and get industry-ready feedback powered by Google Gemini.
# </p>
# """, unsafe_allow_html=True)

# st.write("---")

# # -----------------------
# # 💡 About the Project
# # -----------------------
# with st.expander("💡 About the Project", expanded=True):
#     st.markdown("""
# **AI Resume Reviewer** simplifies your resume improvement process using generative AI:

# - Analyzes your resume content  
# - Identifies strengths & weaknesses  
# - Suggests improvements for clarity, relevance, and professionalism  
# - Gives tips tailored to your target job or industry  
# - No downloads or coding required — upload and get instant feedback  
#     """)

# st.write("---")

# # -----------------------
# # 📤 Upload Resume
# # -----------------------
# st.subheader("📤 Upload Your Resume")

# uploaded_file = st.file_uploader(
#     "Upload your resume (PDF or TXT)",
#     type=["pdf", "txt"]
# )

# # -----------------------
# # 🎯 Job Context Input
# # -----------------------
# st.subheader("🎯 Target Job / Industry")

# job_title = st.text_input("Job Title")
# job_description = st.text_area(
#     "Job Description / Requirements",
#     height=150,
#     placeholder="Paste job description or describe the role you're targeting..."
# )

# st.write("---")

# # -----------------------
# # 🔍 Analyze Button
# # -----------------------
# if st.button("🔍 Analyze My Resume", type="primary"):
#     if not uploaded_file:
#         st.error("Please upload your resume first.")
#     else:
#         with st.spinner("Analyzing your resume with Gemini... ⏳"):

#             # Read uploaded file
#             resume_text = uploaded_file.read().decode("utf-8", errors="ignore")

#             # Prompt to Gemini
#             prompt = f"""
# You are an expert resume reviewer. Analyze the following resume and give clear, structured feedback.

# Resume Content:
# {resume_text}

# Job Target:
# - Job Title: {job_title}
# - Job Description / Requirements: {job_description}

# Please provide analysis in the following format:

# 1. **Overall Summary**
# 2. **Strengths**
# 3. **Weaknesses**
# 4. **Suggestions for Improvement**
# 5. **ATS Optimization Tips**
# 6. **Job Target Alignment Assessment**
# 7. **Rewrite Suggestions (Bullet Points / Summary / Skills)**

# Make the feedback detailed, actionable, and easy to follow.
# """

#             response = model.generate_content(prompt)

#         # Display response
#         st.subheader("📊 Resume Analysis Report")
#         st.markdown(response.text)

#         st.success("Done! Scroll up to view your report.")

# # -----------------------
# # 📌 Footer
# # -----------------------
# st.write("---")
# st.markdown(
#     "<p style='text-align:center; color:gray;'>Built with ❤️ using Streamlit + Google Gemini</p>",
#     unsafe_allow_html=True
# )











# import streamlit as st
# import google.generativeai as genai
# from io import BytesIO
# from PyPDF2 import PdfReader
# from docx import Document
# import json
# import re

# # -----------------------
# # Configure Gemini API
# # -----------------------
# api_key = st.secrets.get("GOOGLE_API_KEY", None)
# if not api_key:
#     st.error("Missing Google API key in st.secrets['GOOGLE_API_KEY']. Add it and redeploy.")
#     st.stop()

# genai.configure(api_key=api_key)
# # model = genai.GenerativeModel("gemini-1.5-flash")  # adjust model if needed
# model = genai.GenerativeModel("gemini-2.5-flash")

# # -----------------------
# # Page config
# # -----------------------
# st.set_page_config(page_title="AI Resume Reviewer", page_icon="📄", layout="wide")

# # -----------------------
# # Helper: extract text from uploaded file
# # -----------------------
# def extract_text_from_file(uploaded_file) -> str:
#     """
#     Accepts a Streamlit UploadedFile and returns extracted text.
#     Supports: .pdf, .docx, .txt
#     """
#     name = uploaded_file.name.lower()
#     uploaded_file.seek(0)
#     data = uploaded_file.read()

#     # Handle PDF
#     if name.endswith(".pdf"):
#         try:
#             reader = PdfReader(BytesIO(data))
#             text_pages = []
#             for page in reader.pages:
#                 page_text = page.extract_text()
#                 if page_text:
#                     text_pages.append(page_text)
#             return "\n\n".join(text_pages).strip()
#         except Exception as e:
#             st.warning(f"PDF parsing error: {e}")
#             return ""

#     # Handle DOCX
#     if name.endswith(".docx"):
#         try:
#             doc = Document(BytesIO(data))
#             paragraphs = [p.text for p in doc.paragraphs if p.text]
#             return "\n\n".join(paragraphs).strip()
#         except Exception as e:
#             st.warning(f"DOCX parsing error: {e}")
#             return ""

#     # Handle plain text
#     if name.endswith(".txt"):
#         try:
#             return data.decode("utf-8", errors="ignore").strip()
#         except Exception:
#             return data.decode("latin-1", errors="ignore").strip()

#     return ""

# # -----------------------
# # Truncate helper (avoid very long prompts)
# # -----------------------
# def truncate_text(text: str, max_chars: int = 30000) -> (str, bool):
#     if len(text) <= max_chars:
#         return text, False
#     truncated = text[:max_chars]
#     return truncated, True

# # -----------------------
# # 🎯 Header Section
# # -----------------------
# st.markdown("""
# <h1 style='text-align:center; margin-bottom:0;'>📄 AI Resume Reviewer</h1>
# <p style='text-align:center; font-size:18px; color:gray;'>
# Upload your resume and get industry-ready feedback powered by Google Gemini.
# </p>
# """, unsafe_allow_html=True)


# # # -----------------------
# # # UI
# # # -----------------------
# # st.title("📄 AI Resume Reviewer")
# # # st.write("Upload your resume and optionally paste a job description. The app extracts text from the file, sends it to Gemini, and returns a structured analysis.")
# # st.write("Upload your resume and get industry-ready feedback powered by Google Gemini.")

# st.write("---")

# # -----------------------
# # 💡 About the Project
# # -----------------------
# with st.expander("💡 About the Project", expanded=True):
#     st.markdown("""
# **AI Resume Reviewer** simplifies your resume improvement process using generative AI:

# - Analyzes your resume content  
# - Identifies strengths & weaknesses  
# - Suggests improvements for clarity, relevance, and professionalism  
# - Gives tips tailored to your target job or industry  
# - No downloads or coding required — upload and get instant feedback  
#     """)

# st.write("---")

# uploaded_file = st.file_uploader("Upload resume (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"])
# job_title = st.text_input("Target Job Title")
# job_description = st.text_area("Job Description / Requirements", height=150)

# st.markdown("---")

# if st.button("🔍 Analyze My Resume"):
#     if not uploaded_file:
#         st.error("Please upload a resume file first.")
#     elif not job_title:
#         st.error("Please enter the target job title.")
#     elif not job_description:
#         st.error("Please enter the job description.")
#     else:
#         with st.spinner("Extracting text from your document..."):
#             resume_text = extract_text_from_file(uploaded_file)

#         if not resume_text:
#             st.error(
#                 "Could not extract readable text from the uploaded file. "
#                 "If it's a scanned PDF/image, please OCR it first or paste the resume text manually."
#             )
#         else:
#             # Truncate if too long and warn user
#             resume_text, was_truncated = truncate_text(resume_text, max_chars=30000)
#             if was_truncated:
#                 st.warning("Resume text was truncated to fit the model prompt (very long resume).")

#             # Build prompt asking for JSON only
#             prompt = f"""
# You are an expert resume reviewer. Analyze the resume and produce a JSON object ONLY (no surrounding text)
# with the following keys: score, summary, strengths, weaknesses, suggestions, ats_tips, alignment.

# - score: a string percentage like "82%"
# - summary: short paragraph
# - strengths: list of short bullet items
# - weaknesses: list of short bullet items
# - suggestions: list of specific actionable changes
# - ats_tips: list of short tips for ATS optimization
# - alignment: short description of how well it matches the provided job

# Resume:
# {resume_text}

# Job target:
# Job Title: {job_title}
# Job Description / Requirements:
# {job_description}

# Return STRICTLY valid JSON. Example:
# {{"score":"80%","summary":"...","strengths":["..."],"weaknesses":["..."],"suggestions":["..."],"ats_tips":["..."],"alignment":"..."}}
# """

#             # Call Gemini
#             with st.spinner("Calling Gemini for analysis..."):
#                 try:
#                     response = model.generate_content(prompt)
#                     raw = response.text.strip()
#                 except Exception as e:
#                     st.error(f"Error calling Gemini: {e}")
#                     st.stop()

#             # Sometimes the model returns text plus JSON; attempt to extract JSON substring
#             json_text = None
#             try:
#                 # naive search for the first JSON object in the response
#                 match = re.search(r"\{[\s\S]*\}\s*$", raw)
#                 if match:
#                     json_text = match.group(0)
#                 else:
#                     # fallback: try to find the first '{' and last '}' and slice
#                     first = raw.find("{")
#                     last = raw.rfind("}")
#                     if first != -1 and last != -1 and last > first:
#                         json_text = raw[first:last+1]
#             except Exception:
#                 json_text = None

#             if not json_text:
#                 st.error("Gemini did not return a JSON object. Raw output shown below for debugging.")
#                 st.code(raw)
#             else:
#                 try:
#                     result = json.loads(json_text)
#                 except Exception as e:
#                     st.error(f"Failed to parse JSON from the model response: {e}")
#                     st.code(json_text)
#                 else:
#                     # Display structured results
#                     st.success("Resume analysis complete!")
#                     st.header("📊 Results Overview")
#                     st.metric("Resume Score", result.get("score", "N/A"))

#                     st.subheader("🧠 Summary")
#                     st.write(result.get("summary", ""))

#                     st.subheader("⭐ Strengths")
#                     for s in result.get("strengths", []):
#                         st.write("• " + s)

#                     st.subheader("⚠️ Weaknesses")
#                     for w in result.get("weaknesses", []):
#                         st.write("• " + w)

#                     st.subheader("🛠 Suggestions")
#                     for sug in result.get("suggestions", []):
#                         st.write("• " + sug)

#                     st.subheader("📑 ATS Tips")
#                     for tip in result.get("ats_tips", []):
#                         st.write("• " + tip)

#                     st.subheader("🎯 Job Alignment")
#                     st.write(result.get("alignment", ""))

#                     # Keep raw response hidden under expander for debugging
#                     with st.expander("Raw model output (for debugging)"):
#                         st.code(raw)

# st.markdown("---")
# st.caption("Built with Streamlit + Google Gemini")












import streamlit as st
import google.generativeai as genai
from io import BytesIO
from PyPDF2 import PdfReader
from docx import Document
import json
import re

# -----------------------
# Configure Gemini API
# -----------------------
api_key = st.secrets.get("GOOGLE_API_KEY", None)
if not api_key:
    st.error("Missing Google API key in st.secrets['GOOGLE_API_KEY']. Add it and redeploy.")
    st.stop()

genai.configure(api_key=api_key)
# Keep your model choice; change if you encounter model-not-found errors
model = genai.GenerativeModel("gemini-2.5-flash")

# -----------------------
# Page config
# -----------------------
st.set_page_config(page_title="AI Resume Reviewer", page_icon="📄", layout="wide")

# -----------------------
# Helper: extract text from uploaded file
# -----------------------
def extract_text_from_file(uploaded_file) -> str:
    """
    Accepts a Streamlit UploadedFile and returns extracted text.
    Supports: .pdf, .docx, .txt
    """
    name = uploaded_file.name.lower()
    uploaded_file.seek(0)
    data = uploaded_file.read()

    # Handle PDF
    if name.endswith(".pdf"):
        try:
            reader = PdfReader(BytesIO(data))
            text_pages = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_pages.append(page_text)
            return "\n\n".join(text_pages).strip()
        except Exception as e:
            st.warning(f"PDF parsing error: {e}")
            return ""

    # Handle DOCX
    if name.endswith(".docx"):
        try:
            doc = Document(BytesIO(data))
            paragraphs = [p.text for p in doc.paragraphs if p.text]
            return "\n\n".join(paragraphs).strip()
        except Exception as e:
            st.warning(f"DOCX parsing error: {e}")
            return ""

    # Handle plain text
    if name.endswith(".txt"):
        try:
            return data.decode("utf-8", errors="ignore").strip()
        except Exception:
            return data.decode("latin-1", errors="ignore").strip()

    return ""

# -----------------------
# Truncate helper (avoid very long prompts)
# -----------------------
def truncate_text(text: str, max_chars: int = 30000) -> (str, bool):
    if len(text) <= max_chars:
        return text, False
    truncated = text[:max_chars]
    return truncated, True

# -----------------------
# Helper: Extract simple skills block from resume text
# -----------------------
def extract_skills_from_text(text: str, max_chars=2000):
    """
    Heuristic: find a 'Skills' or 'Technical Skills' heading and capture the following block.
    Fallback: try to find common skill-like patterns (comma-separated short tokens).
    Returns a list of cleaned skill strings (deduped).
    """
    # Look for headings "Skills", "Technical Skills", "Skills & Tools", etc.
    headings_pattern = re.compile(r"(?im)^(?:skills|technical skills|skills & tools|technical proficiencies)[\s:\-]*\n+([\s\S]{0,800})")
    m = headings_pattern.search(text)
    block = ""
    if m:
        block = m.group(1)
        # stop at blank line followed by an uppercase heading (likely next section)
        split_match = re.split(r"\n\s*\n(?=[A-Z][a-zA-Z ]{1,40}:?)", block)
        if split_match:
            block = split_match[0]

    # If no block found, attempt to find short comma lists that look like skills
    if not block:
        # find lines with many commas or slashes or bullets
        candidates = re.findall(r"^.*(?:,|/|•|-).*$", text, flags=re.MULTILINE)
        if candidates:
            # choose the longest candidate line (likely the skills list)
            block = max(candidates, key=len)
    block = block[:max_chars]

    # Split block into potential skills: bullets, commas, slashes, semi-colons, newlines
    items = re.split(r"[\n•\-\u2022;]+|,|/|;", block)
    cleaned = []
    for it in items:
        s = it.strip()
        # remove long sentences; keep short tokens/phrases
        if not s:
            continue
        # filter out lines that are sentence-like (more than 7 words)
        if len(s.split()) > 7:
            continue
        # remove trailing colons or punctuation
        s = re.sub(r"[:\.\s]+$", "", s)
        s = re.sub(r"^[\-\u2022\s]+", "", s)
        if len(s) > 1:
            cleaned.append(s)
    # Normalize and dedupe while preserving order
    seen = set()
    skills = []
    for c in cleaned:
        key = c.lower()
        if key not in seen:
            seen.add(key)
            skills.append(c)
    return skills[:100]  # limit

def parse_skills_from_string(s: str):
    parts = re.split(r",|/|;|\n", s)
    skills = [p.strip() for p in parts if p.strip()]
    # dedupe
    seen = set(); out=[]
    for sk in skills:
        k = sk.lower()
        if k not in seen:
            seen.add(k); out.append(sk)
    return out

# -----------------------
# 🎯 Header Section
# -----------------------
st.markdown("""
<h1 style='text-align:center; margin-bottom:0;'>📄 AI Resume Reviewer</h1>
<p style='text-align:center; font-size:18px; color:gray;'>
Upload your resume and get industry-ready feedback powered by Google Gemini.
</p>
""", unsafe_allow_html=True)

st.write("---")

# -----------------------
# 💡 About the Project
# -----------------------
with st.expander("💡 About the Project", expanded=True):
    st.markdown("""
**AI Resume Reviewer** simplifies your resume improvement process using generative AI:

- Analyzes your resume content  
- Identifies strengths & weaknesses  
- Suggests improvements for clarity, relevance, and professionalism  
- Gives tips tailored to your target job or industry  
- No downloads or coding required — upload and get instant feedback  
    """)

st.write("---")

uploaded_file = st.file_uploader("Upload resume (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"])
job_title = st.text_input("Target Job Title")
job_description = st.text_area("Job Description / Requirements", height=150)

st.markdown("---")

# -----------------------
# ANALYZE BUTTON (with skills-extraction + editable skills)
# -----------------------
if st.button("🔍 Analyze My Resume"):
    if not uploaded_file:
        st.error("Please upload a resume file first.")
    elif not job_title:
        st.error("Please enter the target job title.")
    elif not job_description:
        st.error("Please enter the job description.")
    else:
        with st.spinner("Extracting text from your document..."):
            resume_text = extract_text_from_file(uploaded_file)

        if not resume_text:
            st.error(
                "Could not extract readable text from the uploaded file. "
                "If it's a scanned PDF/image, please OCR it first or paste the resume text manually."
            )
        else:
            # Truncate if too long and warn user
            resume_text, was_truncated = truncate_text(resume_text, max_chars=30000)
            if was_truncated:
                st.warning("Resume text was truncated to fit the model prompt (very long resume).")

            # ---------------------------
            # Extract skills and allow user to edit them before calling the model
            # ---------------------------
            extracted_skills = extract_skills_from_text(resume_text)
            extracted_str = ", ".join(extracted_skills) if extracted_skills else ""
            st.subheader("🔎 Extracted Skills (editable)")
            st.caption("I attempted to extract a skills list from your resume. Correct or add skills here if anything is missing or incorrect before analysis.")
            skills_edited_str = st.text_area("Edit extracted skills (comma-separated)", value=extracted_str, height=120)

            skills_for_prompt_list = parse_skills_from_string(skills_edited_str)
            skills_for_prompt = ", ".join(skills_for_prompt_list) if skills_for_prompt_list else "None"

            # ---------------------------
            # Build the stronger prompt asking for verified skills + skill_errors
            # ---------------------------
            prompt = f"""
You are an expert resume reviewer. Analyze the resume and produce a JSON object ONLY (no surrounding text).
The JSON must have these keys: score, summary, strengths, weaknesses, suggestions, ats_tips, alignment, skills, skill_errors.

- score: a string percentage like "82%"
- summary: short paragraph
- strengths: list of short bullet items
- weaknesses: list of short bullet items
- suggestions: list of specific actionable changes
- ats_tips: list of short tips for ATS optimization
- alignment: short description of how well it matches the provided job
- skills: verified list of skills the candidate actually has, based strictly on the resume content
- skill_errors: list describing any skills from the provided 'extracted_skills' that are INCORRECT or NOT FOUND in the resume, or any inferred skills the model added (mark each as 'inferred' or 'not_found' accordingly)

Resume:
{resume_text}

User-provided extracted skills (these were automatically extracted from the resume and may be incorrect or partial):
{skills_for_prompt}

Job target:
Job Title: {job_title}
Job Description / Requirements:
{job_description}

IMPORTANT:
1) Only report skills in 'skills' that are actually supported by explicit text in the resume. If you add a skill because it is *clearly implied* (e.g., code samples, library names), mark it as 'inferred' inside skill_errors. If a user_provided skill is not present in the resume, mark it as 'not_found' inside skill_errors.
2) Return STRICTLY valid JSON. Example:
{{"score":"80%","summary":"...","strengths":["..."],"weaknesses":["..."],"suggestions":["..."],"ats_tips":["..."],"alignment":"...","skills":["Python","TensorFlow"],"skill_errors":[{{"skill":"Go","status":"not_found"}}]}}
3) Do not return any commentary outside the JSON.
"""

            # ---------------------------
            # Call Gemini
            # ---------------------------
            with st.spinner("Calling Gemini for skill-verified analysis..."):
                try:
                    response = model.generate_content(prompt)
                    raw = response.text.strip()
                except Exception as e:
                    st.error(f"Error calling Gemini: {e}")
                    st.stop()

            # ---------------------------
            # Extract/parse JSON from the model response (same logic as before)
            # ---------------------------
            json_text = None
            try:
                match = re.search(r"\{[\s\S]*\}\s*$", raw)
                if match:
                    json_text = match.group(0)
                else:
                    first = raw.find("{")
                    last = raw.rfind("}")
                    if first != -1 and last != -1 and last > first:
                        json_text = raw[first:last+1]
            except Exception:
                json_text = None

            if not json_text:
                st.error("Gemini did not return a JSON object. Raw output shown below for debugging.")
                st.code(raw)
            else:
                try:
                    result = json.loads(json_text)
                except Exception as e:
                    st.error(f"Failed to parse JSON from the model response: {e}")
                    st.code(json_text)
                else:
                    # Display structured results
                    st.success("Resume analysis complete!")
                    st.header("📊 Results Overview")
                    st.metric("Resume Score", result.get("score", "N/A"))

                    st.subheader("🧠 Summary")
                    st.write(result.get("summary", ""))

                    st.subheader("⭐ Strengths")
                    for s in result.get("strengths", []):
                        st.write("• " + s)

                    st.subheader("⚠️ Weaknesses")
                    for w in result.get("weaknesses", []):
                        st.write("• " + w)

                    st.subheader("🛠 Suggestions")
                    for sug in result.get("suggestions", []):
                        st.write("• " + sug)

                    st.subheader("📑 ATS Tips")
                    for tip in result.get("ats_tips", []):
                        st.write("• " + tip)

                    st.subheader("🎯 Job Alignment")
                    st.write(result.get("alignment", ""))

                    # New: show verified skills and skill_errors
                    st.subheader("🧾 Verified Skills (from resume)")
                    verified_skills = result.get("skills", [])
                    if verified_skills:
                        st.write(", ".join(verified_skills))
                    else:
                        st.write("No skills returned by the model.")

                    st.subheader("🔍 Skill Extraction Discrepancies")
                    skill_errors = result.get("skill_errors", [])
                    if skill_errors:
                        for err in skill_errors:
                            # err expected to be dicts like {"skill":"Go","status":"not_found" or "inferred"}
                            if isinstance(err, dict):
                                skill_name = err.get("skill", "<unknown>")
                                status = err.get("status", "")
                                note = err.get("note", "")
                                st.write(f"- **{skill_name}** — {status}" + (f": {note}" if note else ""))
                            else:
                                st.write("- " + str(err))
                    else:
                        st.write("No discrepancies reported.")

                    # Keep raw response hidden under expander for debugging
                    with st.expander("Raw model output (for debugging)"):
                        st.code(raw)

st.markdown("---")
st.caption("Built with Streamlit + Google Gemini")





